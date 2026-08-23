# -*- coding: utf-8 -*-
"""Contenido del documento de arquitectura de CreditSmart."""
from build_doc import *   # noqa: F401,F403  (helpers, estilos y objeto `doc`)
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

configure_styles()
build_header_footer()

C = WD_ALIGN_PARAGRAPH.CENTER
J = WD_ALIGN_PARAGRAPH.JUSTIFY

# ============================================================== PORTADA
for text in (u'S20 · Actividad 1 — Evidencia de Aprendizaje 1',
             u'Diseño de Interfaces Web',
             u'CreditSmart — Arquitectura de la Solución'):
    p = doc.add_paragraph(text, style='Heading 1')
    p.alignment = C
    p.paragraph_format.space_after = Pt(2)

para(u'Documento técnico de arquitectura: capas, entidades, rutas, flujos y '
     u'decisiones de diseño de la plataforma de gestión de solicitudes de crédito.',
     align=C, size=11, italic=True, color='6B7280', space_after=18)

for label, hint in ((u'Integrante', u'[ Escriba aquí su nombre completo ]'),
                    (u'Docente', u'[ Escriba aquí el nombre del docente ]'),
                    (u'Curso y NRC', u'[ Ej.: Diseño de Interfaces Web — PREICA0000X000000 ]')):
    para(label, align=C, size=11, bold=True, space_after=2)
    field_placeholder(hint, width_cm=11.0)

para(u'Ingeniería de Software y Datos', align=C, bold=True, space_after=2)
para(u'Facultad de Ingeniería y Ciencias Agropecuarias', align=C, space_after=2)
para(u'Institución Universitaria Digital de Antioquia', align=C, space_after=2)
field_placeholder(u'[ Ciudad y fecha de entrega ]', width_cm=8.0)

page_break()

# ============================================================== TABLA DE CONTENIDO
h1(u'Tabla de contenido')
para(u'Para actualizar el índice: seleccione la tabla, haga clic derecho y elija '
     u'«Actualizar campos» → «Actualizar toda la tabla».', size=9,
     italic=True, color='6B7280')
add_toc()
page_break()

# ============================================================== 1. INTRODUCCION
h1(u'1. Introducción')
para(u'CreditSmart es una plataforma web para consultar, simular y solicitar '
     u'productos de crédito. Este documento describe la arquitectura de la solución '
     u'construida para la actividad: qué capas la componen, qué responsabilidad tiene '
     u'cada una, cómo se comunican entre sí, qué rutas expone al usuario, qué reglas de '
     u'negocio viven en el núcleo y por qué se tomó cada decisión de diseño.')
para(u'La aplicación está escrita en HTML5, CSS3 y JavaScript estándar (módulos ES), '
     u'sin frameworks, sin proceso de compilación y sin dependencias de terceros en '
     u'tiempo de ejecución. Esa restricción —impuesta por el enunciado de la '
     u'actividad— no obligó a renunciar a la organización del código: la solución '
     u'aplica arquitectura hexagonal (puertos y adaptadores), la regla de dependencia '
     u'de Clean Architecture, el patrón MVC en la capa de presentación y los cinco '
     u'principios SOLID, con las interfaces declaradas explícitamente como contratos '
     u'verificables en tiempo de ejecución.')
para(u'El documento está organizado de lo general a lo particular: primero el objetivo '
     u'y el alcance, después la vista de conjunto y las decisiones de arquitectura, '
     u'luego cada capa en detalle (dominio, aplicación, infraestructura y '
     u'presentación), los flujos completos de las tres pantallas, y por último la '
     u'verificación, el cumplimiento de la rúbrica y los entregables.')

callout(u'Cómo leer este documento',
        u'Las figuras numeradas resumen visualmente cada concepto; las tablas '
        u'contienen el detalle exacto (nombres de archivo, firmas de métodos, valores). '
        u'Los recuadros amarillos son espacios editables que debe completar el '
        u'estudiante antes de entregar.')

# ============================================================== 2. OBJETIVO
h1(u'2. Objetivo')
h2(u'2.1 Objetivo general')
para(u'Aplicar los conceptos fundamentales de HTML5, CSS3 y diseño responsive para '
     u'desarrollar las interfaces de usuario de un sistema web de gestión de '
     u'solicitudes de crédito, implementando buenas prácticas de maquetación, '
     u'estructura semántica y diseño profesional, sobre una arquitectura de software '
     u'que separe la interfaz de las reglas del negocio.')

h2(u'2.2 Objetivos específicos')
for t in (u'Construir las tres pantallas requeridas —catálogo, simulador y formulario '
          u'de solicitud— con estructura HTML5 semántica y validación nativa.',
          u'Diseñar una presentación visual profesional con CSS3: paleta coherente, '
          u'Grid y Flexbox, efectos hover, transiciones y sombras.',
          u'Implementar diseño responsive mobile-first verificado en móvil, tablet y '
          u'escritorio.',
          u'Separar la interfaz de las reglas de negocio mediante arquitectura '
          u'hexagonal, de modo que el catálogo estático pueda sustituirse por una API '
          u'real sin modificar el dominio ni la presentación.',
          u'Documentar las decisiones de diseño y verificarlas con pruebas '
          u'automatizadas ejecutables.'):
    bullet(t)

# ============================================================== 3. ALCANCE
h1(u'3. Alcance y requerimientos')
h2(u'3.1 Requerimientos funcionales')
add_table(
    [u'Cód.', u'Requerimiento', u'Dónde se implementa'],
    [[u'RF-01', u'Mostrar el catálogo de productos de crédito con nombre, '
                u'descripción, montos, tasa, plazo y requisitos.',
      u'CatalogView + CatalogController + ListCreditProductsUseCase'],
     [u'RF-02', u'Buscar productos por texto libre (nombre o descripción), '
                u'filtrando de forma incremental al teclear.',
      u'SimulatorView + SearchCreditProductsUseCase + CreditProduct.matchesName()'],
     [u'RF-03', u'Filtrar productos por rango de monto (5 rangos predefinidos).',
      u'GetAmountRangeFiltersUseCase + AmountRange.overlaps()'],
     [u'RF-04', u'Simular un crédito: dado un producto, un monto y un plazo, calcular '
                u'la cuota mensual, el total de intereses, el total a pagar y la tabla '
                u'de amortización mes a mes.',
      u'SimulatorView + SimulateCreditUseCase + CreditSimulationService'],
     [u'RF-05', u'Capturar una solicitud de crédito con datos personales, del '
                u'crédito y laborales (11 campos en 3 secciones).',
      u'ApplicationView + ApplicationController'],
     [u'RF-06', u'Validar la solicitud: formato de cada campo y coherencia con el '
                u'producto elegido (monto y plazo admitidos).',
      u'Applicant · RequestedCredit · EmploymentInfo · CreditApplicationPolicy'],
     [u'RF-07', u'Estimar la cuota mensual y evaluar la capacidad de pago '
                u'(cuota ≤ 40 % del ingreso declarado).',
      u'CreditApplicationPolicy.assessAffordability()'],
     [u'RF-08', u'Radicar la solicitud con número de referencia y persistirla.',
      u'SubmitCreditApplicationUseCase + LocalStorageCreditApplicationRepository'],
     [u'RF-09', u'Navegar entre las tres pantallas sin recargar la página y con '
                u'URL limpias compartibles.',
      u'HistoryRouter + UrlBuilder + NavbarComponent'],
     [u'RF-10', u'Mostrar una pantalla 404 para rutas inexistentes.',
      u'NotFoundView + NotFoundController'],
     [u'RF-11', u'Notificar al usuario el resultado de sus acciones.',
      u'INotifier → ToastNotifier']],
    widths=[1.6, 7.4, 7.0], align_center_cols=(0,))
cap(u'Requerimientos funcionales y su punto de implementación', kind=u'Tabla')

h2(u'3.2 Requerimientos no funcionales')
add_table(
    [u'Atributo', u'Decisión adoptada', u'Evidencia'],
    [[u'Sin dependencias', u'Cero librerías en tiempo de ejecución; cero paso de '
                           u'compilación. Se abre con cualquier servidor estático.',
      u'No existe package.json de runtime; solo módulos ES nativos'],
     [u'Responsive', u'Mobile-first con dos breakpoints (640 px y 1024 px) más un '
                     u'ajuste para pantallas de 320–400 px.', u'assets/css/07-responsive.css'],
     [u'Accesibilidad', u'Etiquetas semánticas, label asociado a cada control, '
                        u'aria-live para notificaciones, foco visible con :focus-visible.',
      u'index.html · BaseView · 03-base.css'],
     [u'Seguridad de plantillas', u'Todo dato interpolado se escapa por defecto; '
                                  u'un único punto escribe innerHTML.',
      u'presentation/shared/Html.js · ViewRenderer.js'],
     [u'Mantenibilidad', u'Regla de dependencia verificable con tres comandos grep; '
                         u'un solo archivo conoce las clases concretas.',
      u'config/dependencies.js'],
     [u'Verificabilidad', u'Tres suites de pruebas ejecutables con Node, sin '
                          u'framework de test.', u'tests/01, tests/02, tests/03'],
     [u'Rendimiento', u'Sin fuentes web, sin imágenes de producto (emojis), '
                      u'1.681 líneas de CSS frente a 70 KB de Tailwind del original.',
      u'assets/css/'],
     [u'Portabilidad', u'La ruta base se autodetecta: funciona en la raíz del dominio '
                       u'o en un subdirectorio.', u'UrlBuilder.detectBasePath()']],
    widths=[3.0, 7.0, 6.0])
cap(u'Requerimientos no funcionales', kind=u'Tabla')

page_break()

# ============================================================== 4. VISION GENERAL
h1(u'4. Visión general de la solución')
para(u'La aplicación es una SPA (Single Page Application) sin framework: el navegador '
     u'carga un único documento HTML —el shell— que aporta el punto de montaje y la '
     u'cascada de estilos; a partir de ahí, un router propio intercepta la navegación '
     u'y cada ruta pinta su vista en el mismo contenedor. Toda la lógica está '
     u'organizada en cuatro capas con una única dirección de dependencia.')

boxes_row([(u'NAVEGADOR', '6B7280'), (u'index.html\n(shell + CSS)', BRAND),
           (u'src/main.js\n(arranque)', VIOLET),
           (u'dependencies.js\n(34 dependencias)', EMERALD),
           (u'HistoryRouter\n(3 rutas + 404)', AMBER)])
cap(u'Cadena de arranque: del documento HTML al router de la aplicación')

figure('chart-capas-dependencia.png',
       u'Las cuatro capas y la regla de dependencia: la infraestructura y la '
       u'presentación dependen del núcleo; el núcleo no depende de nadie', 16.5)

h2(u'4.1 Cifras de la solución')
add_table(
    [u'Métrica', u'Valor', u'Métrica', u'Valor'],
    [[u'Archivos JavaScript', u'79', u'Líneas de JavaScript', u'6.881'],
     [u'Archivos CSS', u'7 (cascada explícita)', u'Líneas de CSS', u'1.681'],
     [u'Capas', u'4 + configuración', u'Contratos (interfaces)', u'12'],
     [u'Entidades', u'2', u'Value objects', u'10'],
     [u'Casos de uso', u'6 (5 consultas, 1 comando)', u'Servicios de dominio', u'2'],
     [u'Vistas', u'6', u'Controladores', u'4 + 1 decorador'],
     [u'Componentes reutilizables', u'4', u'Adaptadores de infraestructura', u'11'],
     [u'Dependencias del contenedor', u'34', u'Suites de pruebas', u'3 (todas en verde)'],
     [u'Aserciones de prueba', u'240', u'Productos de crédito', u'6'],
     [u'Rutas', u'3 + comodín 404', u'Dependencias de runtime', u'0']],
    widths=[4.5, 3.5, 4.5, 3.5], align_center_cols=(1, 3))
cap(u'Dimensiones del sistema', kind=u'Tabla')

figure('chart-capas.png',
       u'Distribución de los 79 archivos JavaScript entre las capas', 11.5)

h2(u'4.2 Responsabilidad de cada capa')
add_table(
    [u'Capa', u'Responde a', u'Contiene', u'Puede importar'],
    [[u'Dominio', u'¿Qué es verdad siempre, sin navegador ni base de datos?',
      u'2 entidades, 10 value objects, 2 servicios, criterios, errores y 7 puertos',
      u'Nada fuera de sí misma'],
     [u'Aplicación', u'¿Qué debe poder hacer el sistema?',
      u'6 casos de uso, DTOs, mappers, Result y 3 puertos',
      u'Solo dominio'],
     [u'Infraestructura', u'¿Con qué tecnología concreta se hace?',
      u'11 adaptadores: memoria, localStorage, Intl, crypto, Date, consola, DOM, '
      u'History API',
      u'Dominio y aplicación (implementa sus puertos)'],
     [u'Presentación', u'¿Cómo se ve y cómo responde a la interacción?',
      u'6 vistas, 4 controladores, 4 componentes, 1 decorador, 3 puertos y '
      u'utilidades de render',
      u'Dominio y aplicación. Nunca infraestructura'],
     [u'Configuración', u'¿Cómo se ensambla todo?',
      u'Contenedor de dependencias, tabla de rutas y parámetros de la aplicación',
      u'Todas (es el único punto que conoce las clases concretas)']],
    widths=[2.6, 3.6, 5.4, 4.4])
cap(u'Responsabilidad y dependencias permitidas por capa', kind=u'Tabla')

page_break()

# ============================================================== 5. DECISIONES
h1(u'5. Decisiones de arquitectura y por qué se tomaron')
para(u'Esta sección responde a la pregunta «¿por qué así y no de otra forma?». Cada '
     u'decisión se presenta con las alternativas que se descartaron y el criterio que '
     u'inclinó la elección.')

h2(u'5.1 Por qué JavaScript vanilla y no un framework')
add_table(
    [u'Alternativa', u'Ventaja', u'Por qué se descartó / se eligió'],
    [[u'React + Vite + Tailwind (el sitio original)',
      u'Ecosistema, componentes, hot reload',
      u'Descartada: el enunciado pide demostrar dominio de HTML5 y CSS3. Con un '
      u'framework, la maquetación semántica y la cascada de estilos quedan ocultas '
      u'tras el bundler.'],
     [u'jQuery + plantillas',
      u'Curva de aprendizaje baja',
      u'Descartada: añade una dependencia sin aportar estructura; el DOM moderno '
      u'cubre todo lo necesario.'],
     [u'JavaScript vanilla con módulos ES',
      u'Cero dependencias, cero build, se abre con cualquier servidor estático; '
      u'todo el código es propio y explicable.',
      u'ELEGIDA. El precio —escribir a mano el router, el render y la inyección de '
      u'dependencias— es justamente lo que la actividad evalúa.']],
    widths=[4.2, 4.4, 7.4])
cap(u'Decisión 1: tecnología de la interfaz', kind=u'Tabla')

h2(u'5.2 Por qué arquitectura hexagonal, Clean Architecture y MVC juntos')
para(u'Los tres patrones resuelven problemas distintos y se complementan; usarlos '
     u'juntos no es redundante:')
bullet(u'define QUÉ necesita el núcleo mediante puertos, y deja que la '
       u'infraestructura decida CÓMO se satisface. Es lo que permite cambiar el '
       u'catálogo estático por una API HTTP tocando un solo archivo.',
       bold_prefix=u'Arquitectura hexagonal (puertos y adaptadores): ')
bullet(u'fija la dirección de las dependencias hacia el centro. Sin '
       u'esa regla, los puertos existirían pero el dominio acabaría importando '
       u'utilidades del navegador.',
       bold_prefix=u'Clean Architecture: ')
bullet(u'organiza únicamente la capa de presentación: la vista pinta, el '
       u'controlador decide, el modelo son los DTOs que entrega el caso de uso.',
       bold_prefix=u'MVC: ')
para(u'El beneficio concreto y medible: las reglas del negocio (rangos de monto, '
     u'plazos admitidos, capacidad de pago, validación de cédula y correo) se prueban '
     u'con Node sin navegador, porque el dominio no conoce el DOM.')

figure('chart-hexagono.png',
       u'Vista hexagonal: seis adaptadores externos entran al núcleo por puertos '
       u'declarados dentro de él', 16.0)

h2(u'5.3 Por qué contratos verificables en tiempo de ejecución')
para(u'JavaScript no tiene la palabra reservada «interface», de modo que un adaptador '
     u'incompleto normalmente fallaría tarde: al hacer clic en el botón que usa el '
     u'método que falta. La solución adoptada declara cada interfaz como un contrato '
     u'explícito y la verifica al arrancar.')
code_block([
    u"// domain/contracts/Contract.js — declaración de una interfaz",
    u"export const IThing = defineContract('IThing', ['doWork', 'reset']);",
    u"",
    u"// El adaptador la extiende e implementa sus métodos",
    u"export class RealThing extends IThing {",
    u"  doWork() { /* ... */ }",
    u"  reset()  { /* ... */ }",
    u"}",
    u"",
    u"// Quien la consume la exige en su constructor",
    u"assertImplements(instance, IThing);   // ContractViolationError si falta un método",
], caption=u'Declaración, implementación y verificación de un contrato')
para(u'Al arrancar, main.js llama a container.eagerResolveAll(), que construye las 34 '
     u'dependencias y ejecuta todas las verificaciones. Consecuencia práctica: un '
     u'contrato roto impide que cargue la página, en lugar de fallar a mitad de una '
     u'navegación. Los métodos no implementados lanzan NotImplementedError al '
     u'invocarse, y assertImplements lanza ContractViolationError con la lista exacta '
     u'de los métodos que faltan.')

h2(u'5.4 Por qué una SPA con History API en lugar de tres archivos HTML')
callout(u'Punto importante para la sustentación',
        u'El enunciado menciona tres archivos: index.html, simulador.html y '
        u'solicitar.html. Esta solución expone las mismas tres direcciones —«/», '
        u'«/simulador» y «/solicitar»— pero las sirve desde un único documento con un '
        u'router propio basado en la History API. Las URL son limpias, compartibles y '
        u'recargables (el .htaccess reescribe hacia index.html); la navegación no '
        u'recarga la página y la navbar, el footer y las tarjetas de producto existen '
        u'una sola vez en el código en lugar de repetirse en tres archivos. En la '
        u'sección 16 se detalla cómo entregar además las tres páginas independientes '
        u'si el docente exige literalmente los tres archivos.',
        fill=WARN_BG, border=AMBER, icon=u'!')
add_table(
    [u'Alternativa', u'Ventaja', u'Coste'],
    [[u'Tres archivos HTML independientes',
      u'Cumple literalmente el enunciado; cero JavaScript de enrutado',
      u'La navbar, el footer y las tarjetas se duplican en tres archivos: tres sitios '
      u'que corregir ante cualquier cambio'],
     [u'SPA con hash (#/simulador)',
      u'No requiere configuración del servidor',
      u'URL menos limpias y peor indexación'],
     [u'SPA con History API (elegida)',
      u'URL limpias, sin duplicación de marcado, transiciones sin recarga, un solo '
      u'punto de montaje',
      u'Requiere reescritura en el servidor (.htaccess incluido) y un router propio '
      u'de ~200 líneas']],
    widths=[4.6, 5.4, 6.0])
cap(u'Decisión 2: estrategia de navegación', kind=u'Tabla')

h2(u'5.5 Por qué CSS semántico en cascada y no utilidades atómicas')
add_table(
    [u'Criterio', u'Tailwind (original)', u'CSS semántico (esta solución)'],
    [[u'Marcado', u'class="bg-white shadow-sm border-b border-gray-200 sticky top-0 z-50"',
      u'class="navbar"'],
     [u'Cambiar el fondo de la navbar', u'Editar la clase en cada página que la repite',
      u'Una línea en 05-components.css'],
     [u'Tamaño', u'70 KB generados', u'1.681 líneas en 7 archivos, sin build'],
     [u'Dependencias', u'PostCSS + Tailwind + configuración', u'Ninguna'],
     [u'Qué demuestra', u'Uso de una librería de utilidades',
      u'Dominio de la cascada, Grid, Flexbox y media queries —lo que evalúa la rúbrica']],
    widths=[3.6, 6.2, 6.2])
cap(u'Decisión 3: estrategia de estilos', kind=u'Tabla')

h2(u'5.6 Resumen de decisiones')
add_table(
    [u'#', u'Decisión', u'Motivo principal'],
    [[u'1', u'JavaScript vanilla con módulos ES', u'La actividad evalúa HTML5 y CSS3, '
      u'no el manejo de un framework'],
     [u'2', u'Arquitectura hexagonal + Clean + MVC', u'Permite probar las reglas de '
      u'negocio sin navegador y sustituir la fuente de datos sin tocar la interfaz'],
     [u'3', u'Contratos verificados al arrancar', u'Convierte un error silencioso en '
      u'un fallo inmediato y localizado'],
     [u'4', u'SPA con History API', u'URL limpias sin duplicar marcado en tres archivos'],
     [u'5', u'CSS semántico en cascada de 7 archivos', u'Cero !important y un único '
      u'lugar por decisión visual'],
     [u'6', u'Value objects en lugar de primitivos', u'Hace imposible construir un '
      u'producto con tasa negativa o rango invertido'],
     [u'7', u'Los casos de uso devuelven Result, no lanzan', u'La presentación nunca '
      u'necesita try/catch y siempre tiene algo que pintar'],
     [u'8', u'Los casos de uso devuelven DTOs, no entidades', u'Impide que una '
      u'plantilla ejecute reglas de negocio'],
     [u'9', u'localStorage con degradación a memoria', u'La aplicación sigue '
      u'funcionando en modo privado o con el almacenamiento bloqueado'],
     [u'10', u'Un único punto que escribe innerHTML', u'Reduce la superficie de XSS a '
      u'un archivo auditable']],
    widths=[1.0, 6.2, 8.8], align_center_cols=(0,))
cap(u'Las diez decisiones de diseño de la solución', kind=u'Tabla')

page_break()

# ============================================================== 6. ESTRUCTURA
h1(u'6. Estructura del proyecto')
code_block([
    u'crediSmart/',
    u'├── index.html                  Shell: punto de montaje + cascada de CSS',
    u'├── manifest.json               Metadatos PWA (nombre, iconos, color de tema)',
    u'├── .htaccess                   Reescritura SPA, MIME, caché y cabeceras',
    u'├── README.md                   Puesta en marcha',
    u'├── assets/css/                 7 archivos en cascada (1.681 líneas)',
    u'│   ├── 01-reset.css            Normalización',
    u'│   ├── 02-tokens.css           Variables: paleta, tipografía, sombras, radios',
    u'│   ├── 03-base.css             Elementos base, foco, spinner, sr-only',
    u'│   ├── 04-layout.css           .page, .container, .section, grids, stacks',
    u'│   ├── 05-components.css       Navbar, botones, tarjetas, formularios, toasts',
    u'│   ├── 06-pages.css            Hero, simulador, formulario, 404',
    u'│   └── 07-responsive.css       Media queries sm / lg / 400px / print',
    u'├── src/',
    u'│   ├── main.js                 Arranque: contenedor, rutas, router',
    u'│   ├── domain/                 NÚCLEO — 23 archivos',
    u'│   │   ├── entities/           CreditProduct · CreditApplication',
    u'│   │   ├── valueobjects/       Money · InterestRate · Term · AmountRange ·',
    u'│   │   │                       ProductTheme · Applicant · RequestedCredit ·',
    u'│   │   │                       EmploymentInfo',
    u'│   │   ├── services/           CreditApplicationPolicy · CreditSimulationService',
    u'│   │   ├── criteria/           ProductSearchCriteria',
    u'│   │   ├── contracts/          7 puertos + Contract.js',
    u'│   │   └── errors/             DomainError · ValidationError · …',
    u'│   ├── application/            CASOS DE USO — 12 archivos',
    u'│   │   ├── usecases/           6 casos de uso',
    u'│   │   ├── dto/ · mappers/     CreditProductDTO · 2 mappers',
    u'│   │   ├── contracts/          IUseCase · ILogger · INotifier',
    u'│   │   └── shared/             Result',
    u'│   ├── infrastructure/         ADAPTADORES — 11 archivos',
    u'│   │   ├── persistence/        Repositorios, datasource y factory',
    u'│   │   ├── routing/            HistoryRouter',
    u'│   │   ├── formatters/         IntlMoneyFormatter',
    u'│   │   ├── identity/ time/     CryptoIdGenerator · SystemClock',
    u'│   │   └── logging/ notification/  ConsoleLogger · ToastNotifier',
    u'│   ├── presentation/           MVC — 22 archivos',
    u'│   │   ├── views/              6 vistas (BaseView + 5)',
    u'│   │   ├── controllers/        4 controladores + BaseController',
    u'│   │   ├── components/         Navbar · Footer · ProductCard · Alert',
    u'│   │   ├── decorators/         DocumentTitleController',
    u'│   │   ├── contracts/          IView · IController · IRouter',
    u'│   │   └── shared/             Html · UrlBuilder · ViewRenderer',
    u'│   └── config/                 ENSAMBLAJE — 5 archivos',
    u'│       ├── dependencies.js     Composition Root: 34 registros',
    u'│       ├── Container.js        Contenedor con detección de ciclos',
    u'│       ├── routes.js           Tabla de rutas (sin imports)',
    u'│       └── AppConfig.js        Parámetros: locale, moneda, log, toasts',
    u'├── tests/                      3 suites ejecutables con Node',
    u'└── docs/                       21 documentos de diseño',
], caption=u'Árbol de carpetas: cada directorio corresponde a una responsabilidad')

# ============================================================== 7. DOMINIO
page_break()
h1(u'7. Capa de dominio: el núcleo del negocio')
para(u'El dominio contiene lo que seguiría siendo verdad sin navegador, sin base de '
     u'datos y sin pantalla. Por eso tiene prohibido usar document, window, fetch, '
     u'localStorage, console, Date.now(), crypto o Intl: para cada una de esas '
     u'necesidades declara un puerto (IClock, IIdGenerator, IMoneyFormatter) que la '
     u'infraestructura satisface.')

h2(u'7.1 Entidades')
add_table(
    [u'Entidad', u'Identidad', u'Composición', u'Comportamiento propio'],
    [[u'CreditProduct', u'id del producto (raíz de agregado del catálogo)',
      u'AmountRange, InterestRate, Term, ProductTheme; ningún campo es un primitivo '
      u'suelto',
      u'matchesName(), matchesAmountRange(), admitsAmount(), admitsTerm(). Congelada '
      u'tras construirse'],
     [u'CreditApplication', u'id generado por IIdGenerator',
      u'Applicant, RequestedCredit, EmploymentInfo, fecha de creación y estado',
      u'submit(), startReview(), referenceNumber (formato CS-XXXXXXXX). No se congela: '
      u'tiene ciclo de vida']],
    widths=[3.2, 3.6, 5.2, 4.6])
cap(u'Las dos entidades del dominio', kind=u'Tabla')

para(u'Toda invariante se verifica en el constructor: un producto sin nombre, con un '
     u'rango que no sea AmountRange o con una tasa que no sea InterestRate no llega a '
     u'existir. Es el patrón «always-valid domain model»: si el objeto existe, sus '
     u'datos son válidos, y ninguna capa posterior necesita volver a comprobarlo.')

h3(u'Ciclo de vida de una solicitud')
boxes_row([(u'DRAFT\n(borrador)', '6B7280'), (u'SUBMITTED\n(radicada)', BRAND),
           (u'UNDER_REVIEW\n(en estudio)', VIOLET)])
cap(u'Estados de CreditApplication. Las transiciones solo ocurren por métodos con '
    u'nombre de negocio —submit(), startReview()— que validan el estado de partida y '
    u'lanzan si la transición es inválida. No existen setters públicos')

h2(u'7.2 Value objects')
para(u'Un value object no tiene identidad: es su valor. Resuelven la «obsesión por los '
     u'primitivos», el error de representar conceptos del negocio con números y '
     u'cadenas sueltas donde nadie avisa si se invierten dos argumentos.')
add_table(
    [u'Value object', u'Encapsula', u'Invariantes que garantiza'],
    [[u'Money', u'Importe y moneda',
      u'Número finito, no negativo, entero; no se operan monedas distintas'],
     [u'InterestRate', u'Tasa efectiva anual en %', u'Número finito entre 0 y 100'],
     [u'Term', u'Plazo en meses', u'Entero positivo'],
     [u'AmountRange', u'Rango [mínimo, máximo]',
      u'Ambos extremos son Money; máximo ≥ mínimo; admite «sin límite»'],
     [u'ProductTheme', u'Emoji y paleta de color',
      u'Icono no vacío; paleta dentro de las 6 admitidas'],
     [u'Applicant', u'Datos personales del solicitante',
      u'Nombre ≥ 5 caracteres, cédula de 6 a 12 dígitos, correo con formato válido, '
      u'teléfono de 7 a 15 dígitos'],
     [u'RequestedCredit', u'Datos del crédito pedido',
      u'Tipo seleccionado, monto > 0, plazo entero > 0, destino ≥ 10 caracteres'],
     [u'EmploymentInfo', u'Datos laborales',
      u'Empresa y cargo no vacíos, ingresos > 0'],
     [u'Installment', u'Una cuota de la tabla de amortización',
      u'Número entero positivo; los cuatro importes son Money; la cuota es exactamente '
      u'interés más capital'],
     [u'AmortizationPlan', u'El plan de pagos completo de una simulación',
      u'Una cuota por cada mes del plazo, numeradas sin huecos; la suma de los abonos a '
      u'capital es igual al capital prestado; la última cuota deja el saldo en cero']],
    widths=[3.0, 4.4, 9.2])
cap(u'Los diez value objects y sus invariantes', kind=u'Tabla')

code_block([
    u'// Sin value objects: los tres errores pasan desapercibidos',
    u'crearProducto(30_000_000, 5_000_000, 18.5, 60);   // min y max invertidos',
    u'crearProducto(5_000_000, 30_000_000, 1850, 60);   // tasa en puntos base',
    u'crearProducto(5_000_000, 30_000_000, 18.5, -60);  // plazo negativo',
    u'',
    u'// Con value objects: los tres son imposibles',
    u'new CreditProduct({',
    u'  amountRange:  new AmountRange(Money.of(5_000_000), Money.of(30_000_000)),',
    u'  interestRate: InterestRate.ofAnnualPercentage(18.5),  // lanza si >100 o <0',
    u'  maxTerm:      Term.ofMonths(60),                      // lanza si <=0',
    u'});',
], caption=u'Por qué el dominio no usa primitivos sueltos')

h2(u'7.3 Servicios de dominio, criterio y errores')
add_table(
    [u'Artefacto', u'Tipo', u'Qué resuelve'],
    [[u'CreditApplicationPolicy', u'Servicio de dominio (sin estado, puro)',
      u'Reglas que cruzan dos entidades: que el monto y el plazo pedidos estén dentro '
      u'de lo que admite el producto, y que la cuota estimada no supere el 40 % del '
      u'ingreso declarado'],
     [u'CreditSimulationService', u'Servicio de dominio (sin estado, puro)',
      u'Amortización por el sistema francés: cruza producto, monto y plazo para '
      u'producir la cuota fija y la tabla mes a mes. Es el único lugar del proyecto '
      u'donde vive la fórmula de la cuota'],
     [u'ProductSearchCriteria', u'Especificación',
      u'Encapsula «qué se considera una coincidencia»: texto en nombre o descripción, '
      u'y solapamiento de rangos de monto'],
     [u'DomainError', u'Error base', u'Todo error del dominio lleva un código y '
      u'detalles, para que la interfaz decida cómo mostrarlo'],
     [u'ValidationError', u'Error de validación',
      u'Agrupa los errores por campo del formulario, que la vista pinta bajo cada '
      u'control'],
     [u'ContractViolationError', u'Error de contrato',
      u'Se lanza al arrancar si un adaptador no implementa todos los métodos de su '
      u'puerto'],
     [u'NotImplementedError', u'Error de contrato',
      u'Lo lanza un método declarado pero no implementado, al invocarse']],
    widths=[3.8, 3.6, 9.2])
cap(u'Servicios, criterios y jerarquía de errores', kind=u'Tabla')

para(u'La cuota mensual se calcula con el sistema de amortización francés, la '
     u'fórmula estándar de cuota fija. Vive en un solo lugar y tiene dos '
     u'consumidores: el simulador y el estudio de capacidad de pago de una '
     u'solicitud. Duplicarla habría permitido que las dos cifras se separaran con '
     u'cualquier ajuste posterior.')
code_block([
    u'// domain/services/CreditSimulationService.js',
    u'//   C = P · i / (1 − (1+i)^−n)',
    u'//   P = capital solicitado   i = tasa mensual   n = número de cuotas',
    u'static monthlyInstallmentAmount(principal, monthlyRate, months) {',
    u'  if (monthlyRate === 0) return principal / months;   // evita dividir por cero',
    u'  return (principal * monthlyRate) / (1 - Math.pow(1 + monthlyRate, -months));',
    u'}',
    u'',
    u'// domain/services/CreditApplicationPolicy.js — la consume, no la repite',
    u'const cuota = CreditSimulationService.monthlyInstallmentAmount(',
    u'  amount.amount, product.interestRate.monthlyFraction, term.months);',
    u'return { affordable: cuota <= techo, estimatedInstallment: Math.round(cuota) };',
], caption=u'La fórmula de la cuota, declarada una vez y consumida dos veces')

para(u'La tabla de amortización se construye sobre esa cuota, mes a mes: el interés '
     u'del período se calcula sobre el saldo vivo y el resto abona capital, de modo '
     u'que el interés decrece y el abono crece. Como los importes se redondean al '
     u'peso, la última cuota abona todo el saldo restante en lugar de una cifra fija: '
     u'absorbe el residuo del redondeo, igual que hace la banca. AmortizationPlan '
     u'verifica al construirse que la suma de los abonos sea exactamente el capital '
     u'prestado y que el saldo final sea cero.')

add_table(
    [u'Simulación verificada', u'Cuota fija', u'Última cuota', u'Total intereses', u'Saldo final'],
    [[u'Libre Inversión · $10.000.000 · 36 meses · 18,5 % E.A.',
      u'$357.000', u'$356.984', u'$2.851.984', u'$0'],
     [u'Vivienda · $300.000.000 · 240 meses · 11,8 % E.A.',
      u'$3.138.761', u'$3.139.123', u'$453.303.002', u'$0']],
    widths=[6.6, 2.6, 2.6, 3.0, 1.8], align_center_cols=(1, 2, 3, 4))
cap(u'Resultados comprobados en la suite de pruebas. El caso de 240 meses es el que '
    u'demuestra que el reparto del redondeo cierra: a veinte años de cuotas, el '
    u'capital sigue cuadrando al peso', kind=u'Tabla')

h2(u'7.4 Los puertos del dominio')
add_table(
    [u'Puerto', u'Métodos', u'Necesidad que declara', u'Adaptador actual'],
    [[u'ICreditProductRepository', u'findAll, findById, findByCriteria, count',
      u'Consultar el catálogo sin saber de dónde viene',
      u'InMemoryCreditProductRepository'],
     [u'ICreditApplicationRepository', u'save, findById, findAll, nextIdentity',
      u'Persistir solicitudes', u'LocalStorageCreditApplicationRepository'],
     [u'IAmountRangeProvider', u'all, byIndex',
      u'Obtener los rangos de monto del filtro', u'StaticAmountRangeProvider'],
     [u'IClock', u'now, timestamp', u'Saber la fecha sin usar Date.now()',
      u'SystemClock'],
     [u'IIdGenerator', u'generate', u'Generar identidades sin usar crypto',
      u'CryptoIdGenerator'],
     [u'IMoneyFormatter', u'format, formatCompact',
      u'Mostrar importes sin conocer Intl ni el locale', u'IntlMoneyFormatter']],
    widths=[4.4, 4.0, 4.6, 3.8])
cap(u'Puertos declarados por el dominio y adaptadores que los satisfacen hoy',
    kind=u'Tabla')

# ============================================================== 8. APLICACION
page_break()
h1(u'8. Capa de aplicación: casos de uso')
para(u'Un caso de uso responde a «el sistema debe poder hacer X». Es una secuencia que '
     u'orquesta al dominio y a los puertos, no un lugar donde vivan reglas: la '
     u'invariante «el monto debe estar dentro del rango del producto» es dominio; la '
     u'secuencia «validar, aplicar la política, guardar, registrar» es aplicación.')

add_table(
    [u'Caso de uso', u'Tipo', u'Dependencias', u'Devuelve'],
    [[u'ListCreditProductsUseCase', u'Consulta', u'productRepository, productMapper',
      u'Result<{ products, total }> — sirve la ruta «/»'],
     [u'SearchCreditProductsUseCase', u'Consulta', u'productRepository, productMapper',
      u'Result<{ products, matched, total, criteria, isFiltered }>'],
     [u'GetAmountRangeFiltersUseCase', u'Consulta', u'amountRangeProvider',
      u'Result<[{ index, label }]> — las 5 opciones del filtro'],
     [u'GetCreditProductNamesUseCase', u'Consulta', u'productRepository',
      u'Result<string[]> — alimenta el <select> del formulario'],
     [u'SimulateCreditUseCase', u'Consulta', u'productRepository, simulationMapper',
      u'Result<SimulationDTO> — cuota, totales y tabla de amortización'],
     [u'SubmitCreditApplicationUseCase', u'Comando',
      u'applicationRepository, productRepository, clock, logger',
      u'Result<{ reference, status, applicantFirstName, affordability }>']],
    widths=[4.6, 2.0, 4.4, 6.0])
cap(u'Los seis casos de uso: cinco consultas y un comando', kind=u'Tabla')

callout(u'El <select> del formulario no duplica el catálogo',
        u'La lista de tipos de crédito del formulario no está escrita a mano: la '
        u'deriva GetCreditProductNamesUseCase del propio catálogo. Añadir un producto '
        u'nuevo actualiza el catálogo, el simulador y el formulario a la vez. Es un '
        u'ejemplo directo de por qué la lógica no vive en la plantilla.')

h2(u'8.1 Result: los casos de uso nunca lanzan hacia fuera')
code_block([
    u'async execute(rawForm = {}) {',
    u'  try {',
    u'    /* ... orquestación ... */',
    u'    return Result.ok({ reference, status, applicantFirstName, affordability });',
    u'  } catch (err) {',
    u'    this.#logger.warn(\'Solicitud rechazada en validación\', { message: err?.message });',
    u'    return Result.fromError(err);      // el error viaja como valor, no como excepción',
    u'  }',
    u'}',
], caption=u'Patrón Result: el controlador siempre recibe algo que pintar')
para(u'Consecuencia en la interfaz: el controlador no necesita try/catch. Consulta '
     u'result.isOk y, si falló, ya tiene el mensaje y los errores por campo listos '
     u'para mostrarse bajo cada control del formulario.')

h2(u'8.2 DTOs y mappers: por qué la vista no recibe entidades')
para(u'Los casos de uso nunca devuelven entidades ni value objects: los pasan por un '
     u'mapper y devuelven objetos planos y congelados. Si la vista recibiera una '
     u'entidad, podría invocar sus métodos y ejecutar reglas de negocio dentro de una '
     u'plantilla —exactamente lo que la arquitectura busca impedir. Los importes '
     u'llegan ya formateados por IMoneyFormatter, de modo que la plantilla solo '
     u'interpola texto.')

# ============================================================== 9. INFRAESTRUCTURA
h1(u'9. Capa de infraestructura: los adaptadores')
add_table(
    [u'Adaptador', u'Puerto que implementa', u'Tecnología', u'Nota de diseño'],
    [[u'InMemoryCreditProductRepository', u'ICreditProductRepository',
      u'Arreglo en memoria construido por CreditProductFactory',
      u'Devuelve copias defensivas: nadie puede mutar el catálogo desde fuera'],
     [u'LocalStorageCreditApplicationRepository', u'ICreditApplicationRepository',
      u'localStorage con clave «creditsmart:applications»',
      u'Si el almacenamiento está bloqueado (modo privado), degrada a memoria en lugar '
      u'de fallar'],
     [u'StaticAmountRangeProvider', u'IAmountRangeProvider', u'Datos estáticos',
      u'Los 5 rangos del filtro'],
     [u'IntlMoneyFormatter', u'IMoneyFormatter',
      u'Intl.NumberFormat(«es-CO», COP, sin decimales)',
      u'Único lugar que conoce el locale'],
     [u'SystemClock', u'IClock', u'Date', u'Sustituible por un reloj fijo en pruebas'],
     [u'CryptoIdGenerator', u'IIdGenerator', u'crypto.randomUUID()',
      u'Degrada a getRandomValues y, si falta, a tiempo + aleatorio'],
     [u'ConsoleLogger', u'ILogger', u'console', u'Respeta el nivel de AppConfig.logLevel'],
     [u'ToastNotifier', u'INotifier', u'DOM (región aria-live)',
      u'Avisos accesibles con cierre automático a los 6 segundos'],
     [u'HistoryRouter', u'IRouter', u'History API + delegación de clics',
      u'Intercepta los enlaces con data-link; restaura el scroll y aplica el ancla'],
     [u'StaticCreditProductDataSource', u'— (fuente de datos)', u'Objetos congelados',
      u'ÚNICO archivo que hay que sustituir para pasar a una API real'],
     [u'CreditProductFactory', u'— (traducción)', u'—',
      u'Capa anticorrupción: convierte los datos crudos en entidades del dominio']],
    widths=[4.8, 3.6, 3.8, 4.8])
cap(u'Los once adaptadores de infraestructura', kind=u'Tabla')

callout(u'Cómo se pasaría de datos estáticos a una API real',
        u'Se crearía HttpCreditProductRepository, que extiende '
        u'ICreditProductRepository y usa fetch; después se cambiaría una línea en '
        u'config/dependencies.js. Ni el dominio, ni los casos de uso, ni las vistas, '
        u'ni el CSS cambiarían: es la prueba de que la inversión de dependencias está '
        u'bien hecha.')

# ============================================================== 10. PRESENTACION
page_break()
h1(u'10. Capa de presentación: MVC sobre el DOM')
add_table(
    [u'Pieza de MVC', u'Aquí', u'Responsabilidad', u'Prohibiciones'],
    [[u'Modelo', u'DTOs que devuelve el caso de uso',
      u'Datos planos, ya formateados', u'No tiene métodos de negocio'],
     [u'Vista', u'6 vistas que extienden BaseView',
      u'Implementar template(viewModel) y traducir eventos del DOM en intenciones',
      u'No consulta datos, no decide navegación, no filtra ni calcula'],
     [u'Controlador', u'4 controladores que extienden BaseController',
      u'Guardar el estado de la interfaz, invocar casos de uso y construir el '
      u'viewModel',
      u'No importa repositorios ni entidades'],
     [u'Componentes', u'Navbar, Footer, ProductCard, Alert',
      u'Funciones puras render(props) → HTML', u'Sin estado ni efectos']],
    widths=[3.0, 4.0, 5.4, 4.6])
cap(u'Reparto de responsabilidades en la capa de presentación', kind=u'Tabla')

h2(u'10.1 Vistas y controladores por pantalla')
add_table(
    [u'Pantalla', u'Vista', u'Controlador', u'Casos de uso que invoca'],
    [[u'Catálogo (/)', u'CatalogView', u'CatalogController', u'ListCreditProducts'],
     [u'Simulador (/simulador)', u'SimulatorView', u'SimulatorController',
      u'GetAmountRangeFilters + SearchCreditProducts'],
     [u'Solicitud (/solicitar)', u'ApplicationView', u'ApplicationController',
      u'GetCreditProductNames + SubmitCreditApplication'],
     [u'404 (cualquier otra)', u'NotFoundView', u'NotFoundController', u'—'],
     [u'Acceso restringido', u'AccessRestrictedView', u'— (pantalla de sistema)', u'—'],
     [u'Base común', u'BaseView', u'BaseController', u'—']],
    widths=[4.0, 3.6, 4.4, 5.0])
cap(u'Correspondencia pantalla · vista · controlador · casos de uso', kind=u'Tabla')

h2(u'10.2 Dos detalles que evitan errores frecuentes')
bullet(u'Todos los listeners se registran con this.on(...) de BaseView, no con '
       u'addEventListener directo. Así destroy() los retira al cambiar de vista. Sin '
       u'esto, cada re-render del simulador dejaría escuchas huérfanas: una fuga de '
       u'memoria por cada tecla pulsada.',
       bold_prefix=u'Limpieza de listeners: ')
bullet(u'present() se usa al entrar en una ruta y sube el scroll al inicio; '
       u'update() se usa en una interacción interna —como teclear en el buscador— y '
       u'no mueve la página. Además, tras cada re-render se restaura el foco con el '
       u'cursor al final, o el usuario escribiría una letra por pulsación.',
       bold_prefix=u'present() frente a update(): ')

h2(u'10.3 Seguridad de plantillas')
para(u'Todo el HTML se construye con la etiqueta de plantilla html`…` de '
     u'shared/Html.js, que escapa por defecto cualquier valor interpolado. La función '
     u'raw() solo se usa con literales escritos en el código fuente, nunca con datos '
     u'del usuario. Un único archivo del proyecto escribe innerHTML —ViewRenderer—, de '
     u'modo que la superficie de riesgo de XSS es un archivo auditable en lugar de '
     u'estar repartida por veinte plantillas.')

# ============================================================== 11. ENDPOINTS
page_break()
h1(u'11. Rutas y superficie de la aplicación')
para(u'La aplicación no consume un backend propio: los datos del catálogo son '
     u'estáticos y la persistencia es local. Por eso su «superficie» tiene dos niveles: '
     u'las rutas que ve el usuario en la barra de direcciones, y la interfaz interna de '
     u'los puertos, que es el punto exacto donde en el futuro entrarían llamadas HTTP.')

h2(u'11.1 Rutas expuestas al usuario')
add_table(
    [u'Ruta', u'Método', u'Pantalla', u'Título del documento', u'Clave del controlador'],
    [[u'/', u'GET', u'Catálogo de créditos', u'CreditSmart — Catálogo', u'catalogController'],
     [u'/simulador', u'GET', u'Simulador de crédito y catálogo filtrable',
      u'CreditSmart — Simulador', u'simulatorController'],
     [u'/solicitar', u'GET', u'Formulario de solicitud', u'CreditSmart — Solicitar',
      u'applicationController'],
     [u'* (comodín)', u'GET', u'Página 404', u'CreditSmart — Página no encontrada',
      u'notFoundController']],
    widths=[2.6, 1.8, 4.0, 4.6, 4.0], align_center_cols=(1,))
cap(u'Tabla de rutas declarada en config/routes.js. El archivo no tiene imports: solo '
    u'constantes y nombres de controlador, lo que evita dependencias circulares',
    kind=u'Tabla')

para(u'El título del documento no lo fija la vista: lo aplica DocumentTitleController, '
     u'un decorador que envuelve a cada controlador. Añadir un comportamiento común a '
     u'todas las rutas —analítica, permisos, título— no exige tocar los cuatro '
     u'controladores.')

h2(u'11.2 Interfaz interna: los métodos de cada puerto')
para(u'Esta es la tabla que sustituye a una lista de endpoints REST: cada método es el '
     u'punto donde una llamada HTTP entraría el día que exista un backend.')
add_table(
    [u'Puerto', u'Método', u'Devuelve', u'Equivalente REST futuro'],
    [[u'ICreditProductRepository', u'findAll()', u'Promise<CreditProduct[]>',
      u'GET /api/productos'],
     [u'', u'findById(id)', u'Promise<CreditProduct | null>', u'GET /api/productos/{id}'],
     [u'', u'findByCriteria(criteria)', u'Promise<CreditProduct[]>',
      u'GET /api/productos?q=&min=&max='],
     [u'', u'count()', u'Promise<number>', u'HEAD /api/productos'],
     [u'ICreditApplicationRepository', u'save(application)',
      u'Promise<CreditApplication>', u'POST /api/solicitudes'],
     [u'', u'findById(id)', u'Promise<CreditApplication | null>',
      u'GET /api/solicitudes/{id}'],
     [u'', u'findAll()', u'Promise<CreditApplication[]>', u'GET /api/solicitudes'],
     [u'', u'nextIdentity()', u'string', u'—  (identidad generada en el cliente)'],
     [u'IAmountRangeProvider', u'all() · byIndex(i)', u'Promise<AmountRange[]>',
      u'GET /api/rangos'],
     [u'IMoneyFormatter', u'format(money) · formatCompact(money)', u'string',
      u'—  (presentación)'],
     [u'IClock', u'now() · timestamp()', u'Date · number', u'—'],
     [u'IIdGenerator', u'generate(prefix?)', u'string', u'—']],
    widths=[4.4, 4.4, 4.2, 4.0])
cap(u'Superficie interna del sistema y su correspondencia con una API REST futura',
    kind=u'Tabla')

h2(u'11.3 Persistencia y configuración')
add_table(
    [u'Elemento', u'Valor', u'Dónde se define'],
    [[u'Clave de almacenamiento', u'creditsmart:applications',
      u'LocalStorageCreditApplicationRepository'],
     [u'Formato del número de radicación', u'CS-XXXXXXXX (8 caracteres del id)',
      u'CreditApplication.referenceNumber'],
     [u'Locale y moneda', u'es-CO · COP · sin decimales', u'AppConfig'],
     [u'Nivel de registro', u'info', u'AppConfig.logLevel'],
     [u'Duración de los avisos', u'6.000 ms', u'AppConfig.toastTimeoutMs'],
     [u'Ruta base de despliegue', u'autodetectada desde document.baseURI',
      u'UrlBuilder.detectBasePath()'],
     [u'Plazos ofrecidos en el formulario', u'12 · 24 · 36 · 48 · 60 meses',
      u'STATIC_TERM_OPTIONS']],
    widths=[5.0, 6.0, 6.0])
cap(u'Parámetros de configuración y persistencia', kind=u'Tabla')

# ============================================================== 12. FLUJOS
page_break()
h1(u'12. Flujos end-to-end')
para(u'Los diagramas siguientes recorren, capa por capa, lo que ocurre en cada '
     u'interacción. El color de la primera columna indica la capa responsable del paso.')

h2(u'12.1 Arranque de la aplicación')
diagram([
    ('C', u'El navegador carga index.html',
     u'Se aplican los 7 archivos CSS en cascada y se muestra el spinner del shell'),
    ('C', u'main.js → bootstrap()',
     u'Localiza #root; si no existe, muestra «No se pudo iniciar la aplicación»'),
    ('C', u'buildContainer({ config, rootElement })',
     u'Registra 34 factorías perezosas: todavía no se construye nada'),
    ('C', u'container.eagerResolveAll()',
     u'Construye las 34 dependencias y verifica cada adaptador contra su puerto con '
     u'assertImplements. Un contrato roto falla AQUÍ, no a mitad de una navegación'),
    ('C', u'Registro de rutas',
     u'Cada entrada de ROUTE_TABLE se envuelve en DocumentTitleController y se '
     u'registra en el router; se fija el controlador de 404'),
    ('I', u'HistoryRouter.start()',
     u'Escucha popstate y delega los clics de los enlaces internos; resuelve la ruta '
     u'actual'),
])
cap(u'Flujo 1: del documento HTML a la primera pantalla pintada')

h2(u'12.2 Cargar el catálogo (ruta «/»)')
diagram([
    ('I', u'HistoryRouter resuelve «/»',
     u'Libera la vista anterior con dispose() y entrega el control al controlador de '
     u'la ruta'),
    ('P', u'CatalogController.handle()',
     u'Invoca el caso de uso; no conoce repositorios ni entidades'),
    ('A', u'ListCreditProductsUseCase.execute()',
     u'Pide findAll() al repositorio, pasa el resultado por el mapper y devuelve '
     u'Result.ok({ products, total })'),
    ('I', u'InMemoryCreditProductRepository',
     u'Devuelve copias de las 6 entidades construidas por CreditProductFactory'),
    ('P', u'Construcción del viewModel',
     u'DTOs planos con los importes ya formateados en pesos colombianos'),
    ('P', u'CatalogView.template() + ViewRenderer',
     u'Genera el HTML —hero, grid de tarjetas, footer— y lo monta en #root; sube el '
     u'scroll al inicio'),
])
cap(u'Flujo 2: carga del catálogo completo')

h2(u'12.3 Buscar en el simulador (búsqueda incremental)')
diagram([
    ('P', u'El usuario teclea «vivienda»',
     u'La vista lee los campos y llama handlers.onSearch({ query, rangeIndex }): '
     u'transmite una INTENCIÓN, no datos procesados'),
    ('P', u'SimulatorController.onSearch()',
     u'Guarda el estado de interfaz —texto, rango y campo con el foco— y vuelve a '
     u'pedir resultados'),
    ('A', u'SearchCreditProductsUseCase.execute()',
     u'Construye ProductSearchCriteria y consulta findByCriteria()'),
    ('D', u'ProductSearchCriteria.isSatisfiedBy()',
     u'Normaliza el texto (sin acentos ni mayúsculas) y compara con nombre y '
     u'descripción; verifica el solapamiento de rangos con AmountRange.overlaps()'),
    ('P', u'update() en lugar de present()',
     u'Se repinta la lista, se restaura el foco con el cursor al final y la página NO '
     u'salta al inicio'),
])
cap(u'Flujo 3: el filtrado ocurre en el dominio, tres capas por debajo del input')

para(u'Ejemplo verificado del filtro por rango «Hasta $5.000.000»: de los seis '
     u'productos sobreviven cuatro —Libre Inversión, Vehículo, Educativo y '
     u'Libranza—, porque sus rangos se solapan con el filtro; Vivienda (desde $20 M) '
     u'y Empresarial (desde $10 M) quedan fuera. La pantalla informa «Mostrando 4 de '
     u'6 productos».')

h2(u'12.4 Simular un crédito')
diagram([
    ('P', u'El usuario teclea el monto',
     u'La vista lee el formulario y llama handlers.onSimulate({ amount, termInMonths }): '
     u'transmite la intención y el campo que tiene el foco, nada más'),
    ('P', u'SimulatorController.onSimulate()',
     u'Actualiza su estado de interfaz y pide una simulación nueva. No calcula nada: '
     u'no hay una sola operación aritmética sobre dinero en el controlador'),
    ('A', u'SimulateCreditUseCase.execute()',
     u'Busca el producto por id, convierte el texto del formulario en Money y Term '
     u'acumulando los errores de forma, y delega el cálculo'),
    ('D', u'CreditSimulationService.simulate()',
     u'Comprueba que el producto admita ese monto y ese plazo, obtiene la tasa mensual '
     u'equivalente y construye la tabla mes a mes con la cuota del sistema francés'),
    ('D', u'AmortizationPlan',
     u'Verifica las cuatro invariantes del plan: una cuota por mes, numeración '
     u'correlativa, capital cuadrado al peso y saldo final en cero'),
    ('A', u'SimulationMapper.toDTO()',
     u'Aplana el plan en un objeto plano y congelado, con los importes ya formateados '
     u'en pesos colombianos por IMoneyFormatter'),
    ('P', u'update() en lugar de present()',
     u'Se repintan la cuota, los totales y la tabla; el foco vuelve al campo que se '
     u'estaba tecleando y la página no salta al inicio'),
])
cap(u'Flujo 4: el cálculo financiero ocurre en el dominio, tres capas por debajo del '
    u'input; la vista solo interpola texto ya formateado')

para(u'Si el monto se sale del rango del producto, el dominio lanza un error por campo '
     u'que el caso de uso devuelve como valor dentro del Result. La vista lo pinta bajo '
     u'el control correspondiente y —decisión deliberada— conserva en pantalla la '
     u'última cuota válida mientras el usuario corrige: borrarla haría parpadear el '
     u'panel en cada tecla intermedia de un número largo.')

page_break()
h2(u'12.5 Radicar una solicitud de crédito')
diagram([
    ('P', u'El usuario envía el formulario',
     u'La vista entrega los 11 campos crudos al controlador, sin interpretarlos'),
    ('A', u'SubmitCreditApplicationUseCase.execute()',
     u'Orquesta la secuencia completa y captura cualquier error para devolverlo como '
     u'Result'),
    ('D', u'Construcción de los value objects',
     u'Applicant, RequestedCredit y EmploymentInfo validan nombre, cédula, correo, '
     u'teléfono, monto, plazo, destino, empresa, cargo e ingresos'),
    ('D', u'CreditApplicationPolicy.assertAdmissible()',
     u'Comprueba que el monto y el plazo estén dentro de lo que admite el producto '
     u'elegido; si no, lanza ValidationError con el mensaje por campo'),
    ('D', u'assessAffordability()',
     u'Estima la cuota con el sistema francés y verifica que no supere el 40 % del '
     u'ingreso declarado'),
    ('D', u'application.submit()',
     u'Transición DRAFT → SUBMITTED y generación de la referencia CS-XXXXXXXX'),
    ('I', u'LocalStorageCreditApplicationRepository.save()',
     u'Persiste la solicitud; si localStorage no está disponible, guarda en memoria'),
    ('P', u'Confirmación en pantalla',
     u'Se muestra el número de radicación y un aviso accesible; si hubo errores, cada '
     u'mensaje aparece bajo su campo'),
])
cap(u'Flujo 5: validación, política de dominio, radicación y persistencia')

# ============================================================== 13. CATALOGO
page_break()
h1(u'13. Catálogo de productos')
add_table(
    [u'', u'Producto', u'Monto mínimo', u'Monto máximo', u'Tasa E.A.', u'Plazo máx.',
     u'Requisitos'],
    [[u'💳', u'Crédito Libre Inversión', u'$1.000.000', u'$30.000.000', u'18,5 %',
      u'60 meses', u'Mayor de 18 años, cédula, extractos bancarios de los últimos 3 meses'],
     [u'🚗', u'Crédito Vehículo', u'$5.000.000', u'$120.000.000', u'14,2 %',
      u'84 meses', u'Mayor de 21 años, licencia de conducción, carta laboral'],
     [u'🏠', u'Crédito Vivienda', u'$20.000.000', u'$500.000.000', u'11,8 %',
      u'240 meses', u'Mayor de 25 años, ingresos demostrables, promesa de compraventa'],
     [u'🎓', u'Crédito Educativo', u'$500.000', u'$50.000.000', u'10,5 %',
      u'72 meses', u'Mayor de 16 años con codeudor, carta de admisión'],
     [u'🏢', u'Crédito Empresarial', u'$10.000.000', u'$1.000.000.000', u'16,0 %',
      u'120 meses', u'Empresa constituida hace 2 años, estados financieros, RUT'],
     [u'🧾', u'Crédito de Libranza', u'$1.000.000', u'$80.000.000', u'13,5 %',
      u'96 meses', u'Empleado o pensionado con convenio de libranza, certificación '
      u'laboral o de pensión, autorización de descuento por nómina']],
    widths=[0.8, 3.1, 2.3, 2.7, 1.5, 1.7, 4.3], align_center_cols=(0, 2, 3, 4, 5))
cap(u'Los seis productos del catálogo, con el emoji y la paleta que identifican a '
    u'cada uno. Los cinco primeros replican el sitio original; «Crédito de '
    u'Libranza» se añadió en este proyecto con la paleta teal', kind=u'Tabla')

figure('chart-tasas.png', u'Comparación de tasas: el crédito educativo es el más '
       u'barato (10,5 %), la libranza le sigue (13,5 %) y el de libre inversión es '
       u'el más costoso (18,5 %)', 16.0)
figure('chart-rangos.png', u'Rangos de monto en escala logarítmica: el filtro del '
       u'simulador selecciona por solapamiento de rangos, no por pertenencia', 16.0)
figure('chart-plazos.png', u'Plazo máximo por producto: desde 5 años en libre '
       u'inversión hasta 20 años en vivienda', 16.0)

h2(u'13.1 Filtros del simulador')
add_table(
    [u'Opción del filtro', u'Rango', u'Productos que devuelve'],
    [[u'Todos los montos', u'sin límite', u'6 productos'],
     [u'Hasta $5.000.000', u'$0 – $5.000.000',
      u'4 — Libre Inversión, Vehículo, Educativo, Libranza'],
     [u'$5.000.001 – $20.000.000', u'$5.000.001 – $20.000.000',
      u'6 — todos se solapan con este tramo'],
     [u'$20.000.001 – $100.000.000', u'$20.000.001 – $100.000.000',
      u'6 — todos se solapan con este tramo'],
     [u'Más de $100.000.000', u'$100.000.001 – sin límite',
      u'3 — Vehículo, Vivienda, Empresarial']],
    widths=[4.6, 4.6, 6.8])
cap(u'Las cinco opciones del filtro por monto y su resultado', kind=u'Tabla')

# ============================================================== 14. INTERFAZ
page_break()
h1(u'14. Interfaz de usuario: HTML5, CSS3 y responsive')
h2(u'14.1 Estructura HTML5 semántica')
add_table(
    [u'Etiqueta', u'Dónde se usa', u'Para qué'],
    [[u'<nav>', u'NavbarComponent', u'Barra de navegación con los tres enlaces y el '
      u'enlace activo resaltado'],
     [u'<header>', u'CatalogView (hero)', u'Encabezado de la página con el título y '
      u'los dos botones de acción'],
     [u'<main>', u'CatalogView, SimulatorView, ApplicationView',
      u'Contenido principal, uno por pantalla'],
     [u'<section>', u'ApplicationView', u'Las tres secciones del formulario, cada una '
      u'con aria-labelledby'],
     [u'<article>', u'ProductCardComponent', u'Cada tarjeta de producto es contenido '
      u'autónomo'],
     [u'<footer>', u'FooterComponent', u'Pie de página con dos variantes de margen'],
     [u'<form>, <label>, <select>', u'SimulatorView, ApplicationView',
      u'Formularios con label asociado por «for» y validación nativa (required, type, '
      u'min)'],
     [u'aria-live, role, aria-label', u'Notificaciones, alertas y spinner',
      u'Accesibilidad: los avisos se anuncian sin robar el foco']],
    widths=[3.4, 5.4, 7.2])
cap(u'Etiquetas semánticas y su uso real en el proyecto', kind=u'Tabla')

para(u'El formulario de solicitud captura once campos agrupados en tres secciones, y '
     u'la vista no repite marcado por campo: los declara como datos y los recorre.')
add_table(
    [u'Sección', u'Campos', u'Validación'],
    [[u'Datos personales',
      u'Nombre completo, cédula, correo electrónico, teléfono',
      u'Nombre ≥ 5 caracteres · cédula 6–12 dígitos · correo con formato · teléfono '
      u'7–15 dígitos'],
     [u'Datos del crédito',
      u'Tipo de crédito, monto solicitado, plazo en meses, destino',
      u'Tipo obligatorio · monto > 0 y dentro del rango del producto · plazo dentro '
      u'del máximo del producto · destino ≥ 10 caracteres'],
     [u'Información laboral', u'Empresa, cargo, ingresos mensuales',
      u'Empresa y cargo no vacíos · ingresos > 0 y suficientes para la cuota estimada']],
    widths=[3.4, 5.4, 7.2])
cap(u'Los once campos del formulario y sus reglas de validación', kind=u'Tabla')

h2(u'14.2 Diseño CSS3 y presentación visual')
figure('chart-css.png', u'Las 1.681 líneas de CSS repartidas en siete archivos que se '
       u'cargan por especificidad creciente', 16.0)
para(u'La cascada explícita tiene una consecuencia medible: el proyecto no necesita un '
     u'solo !important —salvo en las reglas de impresión, donde es la práctica '
     u'habitual—, porque si una regla de componente debe ganar a una de layout, ya '
     u'gana por orden de carga. Y todas las media queries están en un único archivo: '
     u'para saber qué cambia en móvil se lee un archivo, no siete.')
add_table(
    [u'Recurso de CSS3', u'Uso en el proyecto'],
    [[u'Variables CSS (custom properties)',
      u'02-tokens.css concentra la paleta, la tipografía, los radios, las sombras y el '
      u'espaciado. Los componentes solo consumen variables: ningún archivo posterior '
      u'escribe un color literal'],
     [u'Alias semánticos en dos niveles',
      u'--brand, --surface, --text-body apuntan a los colores crudos. Cambiar la marca '
      u'de azul a verde es reasignar una variable, no editar 40 reglas'],
     [u'CSS Grid',
      u'Grid de productos (1 / 2 / 3 columnas), grid del formulario y grid de filtros'],
     [u'Flexbox', u'Navbar, filas de acciones, footer, pie de las tarjetas'],
     [u'Transiciones y hover',
      u'Elevación de las tarjetas, cambio de color en botones y enlaces, anillo de foco'],
     [u'Sombras y radios', u'Escala de sombras (sm, md, lg) y radios coherentes en '
      u'toda la interfaz'],
     [u'Degradados', u'Hero azul y cabecera de cada tarjeta según la paleta de su '
      u'producto'],
     [u'Nomenclatura BEM relajada',
      u'.product-card, .product-card__name, .product-card--compact: la variante '
      u'compacta del simulador se logra añadiendo una clase, sin duplicar marcado']],
    widths=[5.0, 11.0])
cap(u'Recursos de CSS3 aplicados', kind=u'Tabla')

h2(u'14.3 Diseño responsive')
figure('chart-responsive.png', u'Mobile-first: una columna en móvil, dos en tablet y '
       u'tres en escritorio, con el mismo marcado', 16.5)
add_table(
    [u'Breakpoint', u'Qué cambia'],
    [[u'Base (móvil, < 640 px)',
      u'Una columna en el grid de productos y en el formulario; botones a ancho '
      u'completo; la etiqueta «by FinTech Solutions» se oculta'],
     [u'sm — min-width: 640 px',
      u'Dos columnas en productos, formulario y filtros; el título del hero crece; los '
      u'botones se ponen en fila; aparece la etiqueta de la marca'],
     [u'lg — min-width: 1024 px',
      u'Tres columnas en el grid de productos y más aire en el contenedor'],
     [u'max-width: 400 px (añadido propio)',
      u'Compacta la navbar: corrige un desbordamiento real del sitio original entre '
      u'320 y 400 px, donde la marca y los tres enlaces no caben'],
     [u'@media print (añadido propio)',
      u'Oculta navbar, footer, botones y avisos para imprimir el catálogo como ficha '
      u'de producto']],
    widths=[5.0, 11.0])
cap(u'Los cuatro breakpoints y la hoja de impresión', kind=u'Tabla')

h2(u'14.4 Navegación y usabilidad')
bullet(u'la navbar permanece visible al desplazarse, con el enlace de la '
       u'ruta actual resaltado mediante la clase .navlink--active.',
       bold_prefix=u'Navbar sticky: ')
bullet(u'todos los enlaces internos se generan con '
       u'urlBuilder.href(ROUTES.X) y llevan el atributo data-link, de modo que el '
       u'router los intercepta y la navegación no recarga la página.',
       bold_prefix=u'Enlaces consistentes: ')
bullet(u'al entrar en una ruta la página sube al inicio; si la URL '
       u'trae un ancla, se desplaza hasta ella; al filtrar en el simulador no se mueve.',
       bold_prefix=u'Scroll predecible: ')
bullet(u'cada acción produce un aviso accesible que se anuncia en la '
       u'región aria-live y se cierra a los 6 segundos.',
       bold_prefix=u'Retroalimentación: ')
bullet(u'una ruta inexistente muestra una pantalla 404 con la '
       u'dirección solicitada y un botón para volver al catálogo.',
       bold_prefix=u'Errores con salida: ')

h2(u'14.5 Accesibilidad')
add_table(
    [u'Medida', u'Implementación'],
    [[u'Foco visible', u':focus-visible con anillo de 2 px; ningún outline:none sin '
      u'sustituto en todo el proyecto'],
     [u'Etiquetas asociadas', u'Cada control tiene su <label for="…">; los campos '
      u'obligatorios se marcan también en el texto, no solo con color'],
     [u'Avisos anunciados', u'Región aria-live="polite" para las notificaciones y '
      u'role="alert" en los mensajes de error'],
     [u'Contraste', u'Texto sobre superficie con la paleta gray-600/800 sobre blanco'],
     [u'Sin JavaScript', u'Bloque <noscript> que explica el requisito en lugar de '
      u'mostrar una página vacía'],
     [u'Estado de carga', u'Spinner con role="status" y aria-label="Cargando"']],
    widths=[4.6, 11.4])
cap(u'Medidas de accesibilidad', kind=u'Tabla')

# ============================================================== 15. VERIFICACION
page_break()
h1(u'15. Verificación')
h2(u'15.1 La regla de dependencia se comprueba con tres comandos')
code_block([
    u'# 1. El dominio no importa nada de fuera del dominio',
    u'grep -rn "from \'\\.\\./\\.\\./(application|infrastructure|presentation|config)" src/domain/',
    u'',
    u'# 2. La aplicación solo importa del dominio',
    u'grep -rn "from \'\\.\\./\\.\\./(infrastructure|presentation|config)" src/application/',
    u'',
    u'# 3. La presentación nunca importa infraestructura',
    u'grep -rn "infrastructure" src/presentation/',
    u'',
    u'# Las tres deben devolver CERO líneas. Estado actual: cero en las tres.',
], caption=u'Verificación mecánica de la regla de dependencia')

h2(u'15.2 Suites de pruebas')
add_table(
    [u'Suite', u'Qué verifica', u'Comando', u'Resultado'],
    [[u'01-domain-application.mjs',
      u'85 comprobaciones: invariantes de value objects y entidades, servicios de '
      u'dominio, amortización, criterios de búsqueda, los seis casos de uso y Result',
      u'node tests/01-domain-application.mjs', u'TODO OK'],
     [u'02-presentation-render.mjs',
      u'89 comprobaciones: que las vistas generen el HTML esperado, el simulador con '
      u'su tabla y sus errores, el escapado de datos y la detección de dependencias '
      u'circulares del contenedor',
      u'node tests/02-presentation-render.mjs', u'TODO OK'],
     [u'03-boot-jsdom.mjs',
      u'66 comprobaciones: arranque completo con jsdom, resolución de las 34 '
      u'dependencias, navegación entre rutas, simulador y filtros en vivo, foco tras '
      u'el re-render y pantalla 404',
      u'node tests/03-boot-jsdom.mjs', u'TODO OK']],
    widths=[4.6, 6.4, 3.4, 1.6], align_center_cols=(3,))
cap(u'Las tres suites y su resultado en la última ejecución: 240 aserciones en verde',
    kind=u'Tabla')

para(u'Ejemplos de comprobaciones que imprime la ejecución: «monto fuera de rango '
     u'rechazado por CreditApplicationPolicy» con el mensaje «El monto debe estar entre '
     u'$5.000.000 y $120.000.000 para Crédito Vehículo»; «cuota fija esperada $357.000»; '
     u'«la suma de los abonos a capital es exactamente el capital prestado»; «a 240 '
     u'meses el saldo cierra en cero»; «recalcula al teclear el monto ($ 912.498)»; '
     u'«búsqueda "vivienda" → 1 tarjeta»; «foco conservado tras re-render»; '
     u'«contenedor resuelto (34 dependencias)». La tercera suite requiere jsdom: '
     u'npm install jsdom --no-save.')

h2(u'15.3 Cómo ejecutar el proyecto')
code_block([
    u'# Opción A — Laragon / Apache (el proyecto ya está en C:\\laragon\\www\\crediSmart)',
    u'http://crediSmart.test/',
    u'http://localhost/crediSmart/',
    u'',
    u'# Opción B — cualquier servidor estático',
    u'npx serve .            # o:  python -m http.server 8080',
    u'',
    u'# Requiere HTTP: los módulos ES están bloqueados sobre file:// por CORS.',
    u'# El .htaccess incluido reescribe /simulador y /solicitar hacia index.html.',
], caption=u'Puesta en marcha')

# ============================================================== 16. RUBRICA
page_break()
h1(u'16. Cumplimiento de la rúbrica')
add_table(
    [u'Criterio', u'Pts', u'Evidencia en la solución', u'Autoevaluación'],
    [[u'Estructura HTML5 semántica', u'20',
      u'nav, header, main, section, article, footer y form con label asociado; '
      u'validación nativa (required, type, min); código comentado por capas',
      u'[     ]'],
     [u'Diseño CSS3 y presentación visual', u'20',
      u'Paleta en variables (02-tokens.css), Grid y Flexbox, hover, transiciones, '
      u'sombras y degradados; 1.681 líneas organizadas en 7 archivos sin un solo '
      u'!important',
      u'[     ]'],
     [u'Responsive design', u'10',
      u'Mobile-first con breakpoints en 640 px y 1024 px, ajuste extra para 320–400 px '
      u'y hoja de impresión; grid de 1 / 2 / 3 columnas',
      u'[     ]'],
     [u'Navegación y usabilidad', u'10',
      u'Navbar sticky con enlace activo resaltado, navegación sin recarga, scroll '
      u'predecible, avisos accesibles y pantalla 404 con salida',
      u'[     ]'],
     [u'Contenido de las páginas', u'10',
      u'Las tres pantallas completas; formulario de 11 campos en 3 secciones; '
      u'catálogo con 6 productos detallados, uno de ellos añadido en este proyecto',
      u'[     ]'],
     [u'Repositorio Git y documentación', u'10',
      u'README con puesta en marcha y 21 documentos de diseño en docs/; los espacios '
      u'de la sección 18 recogen la URL, los commits y las capturas',
      u'[     ]'],
     [u'Sustentación sincrónica', u'20',
      u'Este documento sirve de guion: decisiones (sección 5), flujos (sección 12) y '
      u'verificación (sección 15)',
      u'[     ]']],
    widths=[4.0, 1.0, 8.4, 2.6], align_center_cols=(1, 3))
cap(u'Correspondencia entre la rúbrica y la evidencia disponible', kind=u'Tabla')

para(u'Espacio para las observaciones del estudiante o del docente sobre la '
     u'autoevaluación:', space_after=4)
field_placeholder(u'[ Escriba aquí sus observaciones sobre el cumplimiento de cada '
                  u'criterio ]', width_cm=16.0, height_cm=2.5)

# ============================================================== 17. BRECHAS
h1(u'17. Consideraciones frente al enunciado')
para(u'Dos puntos del enunciado exigían una decisión explícita. El primero ya está '
     u'resuelto en el código; el segundo se resuelve con sustentación, no con '
     u'código. Ambos se documentan aquí para que la revisión no dependa de '
     u'suposiciones.')

callout(u'1. RESUELTO — el catálogo ya tiene los 6 productos que pide la rúbrica',
        u'La rúbrica otorga los 10 puntos de «Contenido de las páginas» con «mínimo '
        u'6 productos crediticios con información detallada». El catálogo replicaba '
        u'los 5 del sitio original, así que se añadió «Crédito de Libranza»: '
        u'$1.000.000 – $80.000.000, 13,5 % efectivo anual, hasta 96 meses, emoji 🧾 y '
        u'paleta teal. Es el segundo producto más económico del portafolio, con '
        u'descuento directo de nómina o pensión. La sección 13 lo detalla y los tres '
        u'gráficos ya lo incluyen.',
        fill='ECFDF5', border=EMERALD, icon=u'✓')

para(u'Lo que costó el cambio es la mejor evidencia de que la arquitectura hace lo que '
     u'promete: tres archivos tocados y ninguna vista modificada.')
add_table(
    [u'Archivo', u'Cambio', u'Por qué'],
    [[u'StaticCreditProductDataSource.js',
      u'Una entrada nueva con monto, tasa, plazo, requisitos, emoji y paleta',
      u'Es la única fuente de verdad del catálogo: de ahí se derivan las tarjetas, el '
      u'filtro del simulador y el <select> del formulario'],
     [u'ProductTheme.js (dominio)', u'«teal» añadido a THEME_PALETTES',
      u'El value object solo admite paletas declaradas; sin ese paso el producto no se '
      u'habría podido construir — la invariante hizo su trabajo'],
     [u'02-tokens.css', u'Colores teal y bloque .theme-teal',
      u'Las cinco paletas del original ya estaban en uso; el sexto producto merece una '
      u'identidad visual propia'],
     [u'tests/01, 02 y 03',
      u'Conteos esperados: 6 productos, 4 resultados en el filtro «Hasta $5.000.000» y '
      u'7 opciones en el <select>',
      u'Las tres suites vuelven a imprimir TODO OK']],
    widths=[4.4, 5.6, 6.0])
cap(u'Todo lo que hubo que tocar para añadir un producto al catálogo', kind=u'Tabla')

para(u'Ninguna vista, ningún controlador, ningún caso de uso y ningún componente '
     u'cambiaron: el catálogo, el simulador y el formulario mostraron el producto nuevo '
     u'sin que nadie lo escribiera dos veces.')

callout(u'2. DECISIÓN TOMADA — las tres páginas se entregan como SPA y se sustentan '
        u'como tal',
        u'El enunciado nombra index.html, simulador.html y solicitar.html. La solución '
        u'expone las tres direcciones equivalentes —«/», «/simulador» y «/solicitar»— '
        u'desde un único documento con enrutado propio sobre la History API, y esa es '
        u'la decisión que se defiende en la sustentación. Los argumentos: las URL son '
        u'limpias, compartibles y recargables (el .htaccess reescribe hacia '
        u'index.html); la navbar, el footer y la tarjeta de producto existen una sola '
        u'vez en el código en lugar de repetirse en tres archivos; cambiar de pantalla '
        u'no recarga la página; y las tres vistas comparten la misma cascada de CSS y '
        u'la misma capa de dominio. La sección 5.4 recoge la comparación con las '
        u'alternativas descartadas y la sección 11.1 la tabla de rutas que demuestra '
        u'la equivalencia una a una.',
        fill=NOTE_BG, border=BRAND, icon=u'i')

add_table(
    [u'Punto', u'Estado y acción', u'Listo'],
    [[u'Sexto producto del catálogo',
      u'HECHO — «Crédito de Libranza» añadido; las tres suites en verde', u'✓'],
     [u'Los tres archivos HTML del enunciado',
      u'DECIDIDO — se entrega como SPA y se sustenta con la sección 5.4', u'✓'],
     [u'Capturas de pantalla en el README',
      u'Tomar las tres capturas y enlazarlas en README.md', u'[            ]'],
     [u'Cinco commits descriptivos como mínimo',
      u'Verificar el historial con git log --oneline', u'[            ]'],
     [u'Sustentación sincrónica',
      u'Preparar el recorrido: decisiones, flujos y demostración en vivo',
      u'[            ]']],
    widths=[5.0, 7.6, 3.4], align_center_cols=(2,))
cap(u'Lista de verificación previa a la entrega', kind=u'Tabla')

# ============================================================== 18. ENTREGABLES
page_break()
h1(u'18. Entregables')
h2(u'18.1 Repositorio')
para(u'Dirección del repositorio Git (GitHub, GitLab o Bitbucket):', space_after=4)
field_placeholder(u'[ Pegue aquí la URL del repositorio ]', width_cm=14.0)

para(u'Estructura mínima exigida y su estado en este proyecto:', space_after=6)
add_table(
    [u'Archivo exigido', u'Estado en el proyecto', u'Observación'],
    [[u'index.html', u'Presente', u'Shell de la aplicación con la cascada de CSS'],
     [u'simulador.html', u'Como ruta /simulador', u'Ver la sección 17, punto 2'],
     [u'solicitar.html', u'Como ruta /solicitar', u'Ver la sección 17, punto 2'],
     [u'styles.css', u'Como 7 archivos en assets/css/',
      u'Cascada explícita; equivale a un único styles.css mejor organizado'],
     [u'README.md', u'Presente', u'Falta añadir las capturas de pantalla']],
    widths=[4.0, 5.0, 7.0])
cap(u'Correspondencia con los archivos exigidos', kind=u'Tabla')

h2(u'18.2 Historial de commits')
para(u'Registre aquí al menos cinco commits descriptivos del desarrollo:', space_after=4)
add_table(
    [u'#', u'Mensaje del commit', u'Fecha'],
    [[u'1', u'[                                                              ]', u'[          ]'],
     [u'2', u'[                                                              ]', u'[          ]'],
     [u'3', u'[                                                              ]', u'[          ]'],
     [u'4', u'[                                                              ]', u'[          ]'],
     [u'5', u'[                                                              ]', u'[          ]'],
     [u'6', u'[                                                              ]', u'[          ]']],
    widths=[1.0, 11.0, 4.0], align_center_cols=(0, 2))
cap(u'Historial de commits del desarrollo', kind=u'Tabla')

h2(u'18.3 Capturas de pantalla')
for titulo, hint in (
        (u'Captura 1 — Página principal (catálogo de créditos)',
         u'[ Inserte aquí la captura del catálogo: hero, grid de tarjetas y footer ]'),
        (u'Captura 2 — Simulador (búsqueda y filtros)',
         u'[ Inserte aquí la captura del simulador con un filtro aplicado y el '
         u'contador de resultados ]'),
        (u'Captura 3 — Formulario de solicitud',
         u'[ Inserte aquí la captura del formulario con sus tres secciones ]'),
        (u'Captura 4 — Vista responsive en móvil',
         u'[ Inserte aquí la captura en ancho de móvil (menos de 640 px) ]')):
    para(titulo, bold=True, size=10.5, space_after=4)
    field_placeholder(hint, width_cm=16.0, height_cm=5.5)

# ============================================================== 19. CONCLUSIONES
page_break()
h1(u'19. Conclusiones')
para(u'La actividad pedía interfaces web con HTML5, CSS3 y diseño responsive. La '
     u'solución cumple ese alcance y añade una organización interna que hace el '
     u'resultado sostenible: las reglas del negocio están en un núcleo que no conoce el '
     u'navegador, y la interfaz consume esas reglas a través de casos de uso que '
     u'devuelven datos planos.')
para(u'Las consecuencias concretas de esa separación son verificables, no teóricas. '
     u'Las reglas de crédito —rangos, plazos, capacidad de pago, validación de cédula '
     u'y correo— se prueban con Node sin abrir un navegador. Sustituir el catálogo '
     u'estático por una API HTTP exige crear un adaptador y cambiar una línea del '
     u'contenedor de dependencias, sin tocar ni el dominio, ni las vistas, ni el CSS. '
     u'Y un adaptador incompleto impide que la página cargue, en lugar de fallar tres '
     u'clics después.')
para(u'El coste de la decisión también es honesto: hubo que escribir a mano el router, '
     u'el mecanismo de render, el contenedor de dependencias y el sistema de contratos '
     u'—en total 79 archivos y 6.881 líneas de JavaScript— donde un framework habría '
     u'resuelto parte de eso. En el contexto de esta actividad ese coste es la '
     u'evidencia del aprendizaje: cada pieza es propia y explicable.')
para(u'Los dos puntos abiertos frente al enunciado quedaron cerrados. El catálogo ya '
     u'ofrece los seis productos que pide la rúbrica: bastó una entrada en el '
     u'datasource, una paleta nueva en el value object y la actualización de los '
     u'conteos esperados en las pruebas —ninguna vista cambió—, lo que confirma en '
     u'la práctica el beneficio que la sección 5 anunciaba en teoría. Y las tres '
     u'páginas se entregan como SPA con sus tres rutas equivalentes, decisión que se '
     u'sustenta con los argumentos de la sección 5.4. Queda únicamente completar el '
     u'README con las capturas y registrar el historial de commits.')

h1(u'20. Referencias')
for ref in (
    u'MDN Web Docs. HTML: Lenguaje de etiquetas de hipertexto — elementos semánticos. '
    u'developer.mozilla.org',
    u'MDN Web Docs. CSS: Hojas de estilo en cascada — Grid Layout, Flexbox y media '
    u'queries. developer.mozilla.org',
    u'MDN Web Docs. History API y el evento popstate. developer.mozilla.org',
    u'W3C. Web Content Accessibility Guidelines (WCAG) 2.1.',
    u'Cockburn, A. Hexagonal Architecture (Ports and Adapters), 2005.',
    u'Martin, R. C. Clean Architecture: A Craftsman’s Guide to Software Structure and '
    u'Design. Prentice Hall, 2017.',
    u'Evans, E. Domain-Driven Design: Tackling Complexity in the Heart of Software. '
    u'Addison-Wesley, 2003.',
    u'Vernon, V. Implementing Domain-Driven Design. Addison-Wesley, 2013.',
    u'Documentación interna del proyecto: docs/master.md — índice de los 21 documentos '
    u'de diseño (arquitectura, entidades, value objects, contratos, casos de uso, '
    u'adaptadores, inyección de dependencias, enrutado, estilos, flujos, pruebas y '
    u'convenciones).',
):
    bullet(ref)

para(u'Nota sobre el origen del contenido funcional: la aplicación es una '
     u'reconstrucción del sitio de referencia sweet-smart-credit-path.base44.app '
     u'(originalmente React + Vite + Tailwind) en HTML, CSS y JavaScript estándar. Se '
     u'replicaron los textos, productos, colores y breakpoints del original; la '
     u'arquitectura, el enrutado, las validaciones y las pruebas son desarrollo propio.',
     size=9, italic=True, color='6B7280')

doc.save(OUT_DOCX)
print(u'Documento generado: %s' % OUT_DOCX)
print(u'Figuras: %d   Tablas numeradas: %d' % (FIG_N[0], TAB_N[0]))
