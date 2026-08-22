# -*- coding: utf-8 -*-
"""
Genera el documento de arquitectura de CreditSmart en formato .docx,
siguiendo el patron de diseno de docs/iudigital_doc/Pedraza_Jeremy_TallerDOFA.docx:

  - Arial 11 pt, margenes de 1 pulgada, tamano carta
  - Titulo 1: Arial 16 pt negrita #1F3864   Titulo 2: 13 pt negrita #2E75B6
  - Titulo 3: 12 pt #1F4D78
  - Tablas centradas, bordes #7F7F7F, fila de encabezado #1F3864 con texto blanco 9 pt
  - Encabezado con el logo institucional, pie con "Pagina X de Y"
  - Portada centrada con los datos academicos
"""
import os
import zipfile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Pt, RGBColor

HERE = os.path.dirname(os.path.abspath(__file__))
IMG = os.path.join(HERE, 'img')
REF_DOCX = r'C:/laragon/www/crediSmart/docs/iudigital_doc/Pedraza_Jeremy_TallerDOFA.docx'
OUT_DIR = r'C:/laragon/www/crediSmart/docs/iudigital_doc'
# NUNCA escribir sobre el archivo que edita el usuario a mano:
# la salida del generador vive en su propio nombre, con sufijo _generado.
OUT_DOCX = os.path.join(OUT_DIR, 'CreditSmart_Arquitectura_de_la_Solucion_generado.docx')

# ----------------------------------------------------------------- paleta
NAVY = '1F3864'        # titulos 1 y encabezados de tabla (patron del documento guia)
BLUE = '2E75B6'        # titulos 2
BLUE3 = '1F4D78'       # titulos 3
BRAND = '1D4ED8'       # azul de marca del proyecto (blue-700)
EMERALD = '10B981'
VIOLET = '8B5CF6'
AMBER = 'F59E0B'
ROSE = 'F43F5E'
GRAY_BORDER = '7F7F7F'
ROW_ALT = 'F2F6FB'
CODE_BG = 'F3F4F6'
NOTE_BG = 'EFF6FF'
WARN_BG = 'FFFBEB'
FIELD_BG = 'FFF2CC'    # espacios por llenar

# Colores por capa, usados en los diagramas de flujo
LAYER_COLORS = {
    'P': (BRAND, u'PRESENTACIÓN'),
    'A': (VIOLET, u'APLICACIÓN'),
    'D': (NAVY, u'DOMINIO'),
    'I': (AMBER, u'INFRAESTRUCTURA'),
    'C': ('6B7280', u'CONFIGURACIÓN'),
}

doc = Document()


# ================================================================= utilidades
def set_cell_shading(cell, hex_fill):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_fill)
    cell._tc.get_or_add_tcPr().append(shd)


def set_cell_margins(cell, top=40, bottom=40, left=80, right=80):
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement('w:tcMar')
    for name, val in (('top', top), ('left', left), ('bottom', bottom), ('right', right)):
        node = OxmlElement('w:' + name)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        mar.append(node)
    tcPr.append(mar)


def set_table_borders(table, color=GRAY_BORDER, sz=4):
    tblPr = table._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement('w:' + edge)
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), str(sz))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        borders.append(el)
    tblPr.append(borders)


USABLE_CM = 16.4          # ancho util: carta menos dos margenes de 1 pulgada


def set_fixed_widths(table, widths):
    """Fija el ancho de cada columna: Word ignora las anchuras si el layout es auto."""
    total = float(sum(widths))
    scaled = [w * USABLE_CM / total for w in widths]
    table.autofit = False
    tblPr = table._tbl.tblPr
    layout = OxmlElement('w:tblLayout')
    layout.set(qn('w:type'), 'fixed')
    tblPr.append(layout)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(int(USABLE_CM * 567)))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)
    grid = table._tbl.find(qn('w:tblGrid'))
    if grid is not None:
        for col, w in zip(grid.findall(qn('w:gridCol')), scaled):
            col.set(qn('w:w'), str(int(w * 567)))
    for row in table.rows:
        for cell, w in zip(row.cells, scaled):
            cell.width = Cm(w)
    return table


def keep_together(table):
    """Evita que una caja (codigo, aviso, hueco editable) se parta entre paginas."""
    for row in table.rows:
        trPr = row._tr.get_or_add_trPr()
        el = OxmlElement('w:cantSplit')
        el.set(qn('w:val'), 'true')
        trPr.append(el)
    return table


def repeat_header_row(row):
    trPr = row._tr.get_or_add_trPr()
    el = OxmlElement('w:tblHeader')
    el.set(qn('w:val'), 'true')
    trPr.append(el)


def cell_text(cell, text, size=9, bold=False, color=None, align=None,
              italic=False, mono=False):
    cell.text = ''
    p = cell.paragraphs[0]
    pf = p.paragraph_format
    pf.space_before = Pt(1)
    pf.space_after = Pt(1)
    if align is not None:
        p.alignment = align
    first = True
    for line in str(text).split('\n'):
        if not first:
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(1)
            if align is not None:
                p.alignment = align
        first = False
        run = p.add_run(line)
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        if mono:
            run.font.name = 'Consolas'
        if color:
            run.font.color.rgb = RGBColor.from_string(color)
    return cell


def add_table(headers, rows, widths=None, header_fill=NAVY, zebra=True,
              font_size=9, align_center_cols=()):
    """Tabla con el formato del documento guia."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table)

    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        set_cell_shading(c, header_fill)
        set_cell_margins(c)
        cell_text(c, h, size=font_size, bold=True, color='FFFFFF')
    repeat_header_row(table.rows[0])

    for r, data in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(data):
            c = cells[i]
            set_cell_margins(c)
            if zebra and r % 2 == 1:
                set_cell_shading(c, ROW_ALT)
            align = WD_ALIGN_PARAGRAPH.CENTER if i in align_center_cols else None
            mono = str(value).startswith('`')
            text = str(value).strip('`') if mono else value
            cell_text(c, text, size=font_size, align=align, mono=mono)

    if widths:
        set_fixed_widths(table, widths)
    doc.add_paragraph()
    return table


def h1(text):
    p = doc.add_paragraph(text, style='Heading 1')
    return p


def h2(text):
    return doc.add_paragraph(text, style='Heading 2')


def h3(text):
    return doc.add_paragraph(text, style='Heading 3')


def para(text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=11, bold=False,
         italic=False, color=None, space_after=8):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return p


def rich(parts, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=11, space_after=8):
    """parts = [(texto, 'b'|'i'|'code'|''), ...]"""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    for text, fmt in parts:
        run = p.add_run(text)
        run.font.size = Pt(size if fmt != 'code' else size - 1.5)
        run.bold = 'b' in fmt
        run.italic = 'i' in fmt
        if fmt == 'code':
            run.font.name = 'Consolas'
            run.font.color.rgb = RGBColor.from_string('1F4D78')
    return p


def bullet(text, level=0, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8 + 0.6 * level)
    p.paragraph_format.space_after = Pt(3)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    mark = p.add_run(('•  ' if level == 0 else '–  '))
    mark.font.size = Pt(11)
    mark.font.color.rgb = RGBColor.from_string(BRAND)
    mark.bold = True
    if bold_prefix:
        b = p.add_run(bold_prefix)
        b.font.size = Pt(11)
        b.bold = True
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p


def code_block(lines, caption=None):
    """Bloque de codigo: tabla de una celda con fondo gris y fuente monoespaciada."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color='D1D5DB', sz=4)
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, CODE_BG)
    set_cell_margins(cell, top=80, bottom=80, left=140, right=140)
    cell.text = ''
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor.from_string('1F2937')
    keep_together(table)
    if caption:
        cap(caption)
    else:
        doc.add_paragraph()
    return table


def callout(title, text, fill=NOTE_BG, border=BRAND, icon='i'):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color=border, sz=8)
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=100, bottom=100, left=140, right=140)
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(u'%s  %s' % (icon, title))
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string(border)
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r2 = p2.add_run(text)
    r2.font.size = Pt(9.5)
    keep_together(table)
    doc.add_paragraph()
    return table


def figure(filename, caption, width_cm=16.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(os.path.join(IMG, filename), width=Cm(width_cm))
    cap(caption)


FIG_N = [0]
TAB_N = [0]


def cap(text, kind='Figura'):
    if kind == 'Figura':
        FIG_N[0] += 1
        n = FIG_N[0]
    else:
        TAB_N[0] += 1
        n = TAB_N[0]
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(u'%s %d. %s' % (kind, n, text))
    r.font.size = Pt(8.5)
    r.italic = True
    r.font.color.rgb = RGBColor.from_string('6B7280')
    return p


def field_placeholder(hint, width_cm=10.0, height_cm=None):
    """Espacio editable: celda amarilla clara con la indicacion en cursiva."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color=AMBER, sz=4)
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, FIELD_BG)
    set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
    cell_text(cell, hint, size=9, italic=True, color='8A6100',
              align=WD_ALIGN_PARAGRAPH.CENTER)
    if height_cm:
        tr = table.rows[0]._tr
        trPr = tr.get_or_add_trPr()
        h = OxmlElement('w:trHeight')
        h.set(qn('w:val'), str(int(height_cm * 567)))
        h.set(qn('w:hRule'), 'atLeast')
        trPr.append(h)
    table.rows[0].cells[0].width = Cm(width_cm)
    keep_together(table)
    doc.add_paragraph()
    return table


def diagram(rows, headers=(u'Capa', u'Paso', u'Qué ocurre')):
    """Diagrama de flujo como tabla: la primera columna es un chip de color por capa."""
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        set_cell_shading(c, NAVY)
        set_cell_margins(c)
        cell_text(c, h, size=9, bold=True, color='FFFFFF')
    repeat_header_row(table.rows[0])

    for layer, step, detail in rows:
        cells = table.add_row().cells
        color, name = LAYER_COLORS[layer]
        set_cell_shading(cells[0], color)
        set_cell_margins(cells[0])
        cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell_text(cells[0], name, size=6.8, bold=True, color='FFFFFF',
                  align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_margins(cells[1])
        cell_text(cells[1], step, size=8.5, bold=True, color='1F2937')
        set_cell_margins(cells[2])
        cell_text(cells[2], detail, size=8.5)

    set_fixed_widths(table, [3.1, 4.8, 8.5])
    doc.add_paragraph()
    return table


def boxes_row(items, fill_map=None):
    """Fila de cajas de colores con flechas, para diagramas simples."""
    cols = len(items) * 2 - 1
    table = doc.add_table(rows=1, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tblPr = table._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement('w:' + edge)
        el.set(qn('w:val'), 'none')
        borders.append(el)
    tblPr.append(borders)
    n_boxes = len(items)
    arrow_w = 0.7
    box_w = (USABLE_CM - arrow_w * (n_boxes - 1)) / n_boxes
    widths = []
    idx = 0
    for i, (label, color) in enumerate(items):
        if i:
            c = table.rows[0].cells[idx]
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell_text(c, u'▶', size=11, bold=True, color='6B7280',
                      align=WD_ALIGN_PARAGRAPH.CENTER)
            widths.append(arrow_w)
            idx += 1
        c = table.rows[0].cells[idx]
        set_cell_shading(c, color)
        set_cell_margins(c, top=90, bottom=90, left=40, right=40)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell_text(c, label, size=7.5, bold=True, color='FFFFFF',
                  align=WD_ALIGN_PARAGRAPH.CENTER)
        widths.append(box_w)
        idx += 1
    set_fixed_widths(table, widths)
    doc.add_paragraph()
    return table


def page_break():
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


# ================================================================= estilos base
def configure_styles():
    normal = doc.styles['Normal']
    normal.font.name = 'Arial'
    normal.font.size = Pt(11)
    rpr = normal.element.get_or_add_rPr().get_or_add_rFonts()
    for attr in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
        rpr.set(qn(attr), 'Arial')
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    specs = [('Heading 1', 16, NAVY, True, 14, 8),
             ('Heading 2', 13, BLUE, True, 12, 6),
             ('Heading 3', 12, BLUE3, True, 10, 4)]
    for name, size, color, bold, before, after in specs:
        st = doc.styles[name]
        st.font.name = 'Arial'
        st.font.size = Pt(size)
        st.font.bold = bold
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True
        r = st.element.get_or_add_rPr().get_or_add_rFonts()
        for attr in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
            r.set(qn(attr), 'Arial')

    section = doc.sections[0]
    section.page_width = Emu(7772400)     # carta
    section.page_height = Emu(10058400)
    for m in ('left_margin', 'right_margin', 'bottom_margin'):
        setattr(section, m, Emu(914400))   # 1 pulgada
    # El logo del encabezado mide 1,51 cm: el margen superior lo deja respirar
    section.top_margin = Cm(3.3)
    section.header_distance = Cm(1.1)
    section.footer_distance = Emu(457200)


def build_header_footer():
    """Logo institucional a la derecha del encabezado y 'Pagina X de Y' en el pie."""
    section = doc.sections[0]

    logo_path = os.path.join(HERE, 'refmedia', 'word', 'media', 'image1.png')
    if not os.path.exists(logo_path):
        with zipfile.ZipFile(REF_DOCX) as z:
            z.extract('word/media/image1.png', os.path.join(HERE, 'refmedia'))
    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.add_run().add_picture(logo_path, width=Emu(1971675), height=Emu(552450))

    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run(u'Página ')
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string('6B7280')

    def field(instr):
        run = fp.add_run()
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string('6B7280')
        begin = OxmlElement('w:fldChar')
        begin.set(qn('w:fldCharType'), 'begin')
        it = OxmlElement('w:instrText')
        it.set(qn('xml:space'), 'preserve')
        it.text = instr
        end = OxmlElement('w:fldChar')
        end.set(qn('w:fldCharType'), 'end')
        run._r.append(begin)
        run._r.append(it)
        run._r.append(end)

    field('PAGE')
    r = fp.add_run(u' de ')
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string('6B7280')
    field('NUMPAGES')


def add_toc():
    p = doc.add_paragraph()
    run = p.add_run()
    begin = OxmlElement('w:fldChar')
    begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = r'TOC \o "1-3" \h \z \u'
    sep = OxmlElement('w:fldChar')
    sep.set(qn('w:fldCharType'), 'separate')
    placeholder = OxmlElement('w:t')
    placeholder.text = (u'Haga clic derecho aquí y elija «Actualizar campos» '
                        u'para generar la tabla de contenido.')
    end = OxmlElement('w:fldChar')
    end.set(qn('w:fldCharType'), 'end')
    for el in (begin, instr, sep, placeholder, end):
        run._r.append(el)
